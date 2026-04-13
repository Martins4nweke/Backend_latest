from io import BytesIO
from docx import Document

def _add_key_value_table(doc, heading, data):
    if not data:
        return
    doc.add_heading(heading, level=1)
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    for k, v in data.items():
        r = t.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)

def _add_list_of_dicts(doc, heading, rows):
    if not rows:
        return
    doc.add_heading(heading, level=1)
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = '' if row.get(h) is None else str(row.get(h))

def build_word_report(payload):
    doc = Document()
    doc.add_heading('NCD Cascade Allocation Policy Report v6.3', level=0)
    doc.add_paragraph('This report was generated from the evidence-aware allocation backend.')
    if payload.get('executive_summary'):
        doc.add_heading('Executive summary', level=1)
        doc.add_paragraph(payload['executive_summary'])
    _add_key_value_table(doc, 'Key performance indicators', payload.get('kpis', {}))
    _add_key_value_table(doc, 'Budget diagnostics', payload.get('budget_diagnostics', {}))
    _add_key_value_table(doc, 'Structure diagnostics', payload.get('structure_diagnostics', {}))
    _add_key_value_table(doc, 'Equity diagnostics', payload.get('equity_diagnostics', {}))
    _add_key_value_table(doc, 'Evidence retrieval summary', payload.get('evidence_retrieval_summary', {}))
    _add_key_value_table(doc, 'Parameter provenance summary', payload.get('parameter_provenance_summary', {}))
    for title, key in [
        ('Policy advisory brief', 'policy_advisory_brief'),
        ('Operational brief', 'policy_operational_brief'),
        ('Equity brief', 'policy_equity_brief'),
        ('Budget utilisation brief', 'policy_budget_brief'),
        ('Scenario brief', 'policy_scenario_brief'),
        ('Confidence note', 'confidence_note'),
        ('Limitations', 'limitations_note'),
    ]:
        if payload.get(key):
            doc.add_heading(title, level=1)
            doc.add_paragraph(payload.get(key, ''))
    doc.add_heading('Equity-sensitive recommendations', level=1)
    for rec in payload.get('equity_recommendations', []):
        doc.add_paragraph(rec, style='List Bullet')
    province_briefs = payload.get('province_briefs', {})
    if province_briefs:
        doc.add_heading('Province-specific briefs', level=1)
        for province, brief in province_briefs.items():
            doc.add_heading(province, level=2)
            doc.add_paragraph(brief)
    _add_list_of_dicts(doc, 'Province comparative table', payload.get('province_comparative_table', []))
    _add_list_of_dicts(doc, 'Scenario results table', payload.get('scenario_tradeoff_table', []))
    _add_list_of_dicts(doc, 'Parameter provenance preview', payload.get('parameter_provenance_preview', []))
    top_rows = payload.get('grouped_summaries', {}).get('top_rows', [])
    if top_rows:
        show = []
        for row in top_rows:
            show.append({
                'province': row.get('province'),
                'stratum': row.get('stratum_code'),
                'intervention': row.get('intervention_name'),
                'units': row.get('units_allocated'),
                'spend_zar': row.get('spend_zar'),
                'dalys_averted': row.get('dalys_averted'),
                'source_tier': row.get('source_tier'),
            })
        _add_list_of_dicts(doc, 'Top allocation rows', show)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
